from __future__ import annotations

import io
import logging
import zipfile
from pathlib import Path

from interview_app.cv.config import ALLOWED_CV_EXTENSIONS, DEFAULT_CV_MAX_FILE_BYTES
from interview_app.cv.exceptions import CVExtractionError, CVFileValidationError

logger = logging.getLogger("interview_app.cv.document_parser")


def _normalize_extension(filename: str) -> str:
    return Path(filename or "").suffix.lower()


def _is_probably_pdf(data: bytes) -> bool:
    return len(data) >= 4 and data[:4] == b"%PDF"


def _is_probably_docx(data: bytes) -> bool:
    if len(data) < 4 or data[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = zf.namelist()
    except zipfile.BadZipFile:
        return False
    return any(n == "word/document.xml" for n in names)


def validate_cv_upload(
    *,
    filename: str,
    data: bytes,
    max_bytes: int = DEFAULT_CV_MAX_FILE_BYTES,
) -> str:
    """
    Validate filename, size, and magic bytes. Returns normalized extension (e.g. '.pdf').

    Raises:
        CVFileValidationError: When validation fails.
    """
    ext = _normalize_extension(filename)
    if ext not in ALLOWED_CV_EXTENSIONS:
        raise CVFileValidationError(
            "Unsupported file type. Please upload a PDF or DOCX file."
        )
    if not data:
        raise CVFileValidationError("The uploaded file is empty.")
    if len(data) > max_bytes:
        raise CVFileValidationError(
            f"File is too large (max {max_bytes // (1024 * 1024)} MB)."
        )

    if ext == ".pdf" and not _is_probably_pdf(data):
        raise CVFileValidationError("This file does not look like a valid PDF.")
    if ext == ".docx" and not _is_probably_docx(data):
        raise CVFileValidationError("This file does not look like a valid DOCX document.")

    return ext


def extract_text_pdf(data: bytes) -> str:
    """Extract plain text from PDF bytes."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise CVExtractionError(
            "PDF text extraction needs the 'pypdf' package. "
            "From your project folder run: pip install pypdf  "
            "(or pip install -r requirements.txt)."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        logger.warning("PdfReader failed to open PDF: %s", type(exc).__name__)
        raise CVExtractionError("Could not read this PDF (file may be corrupted).") from exc

    n_pages = len(reader.pages)
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t)
    joined = "\n\n".join(parts).strip()
    if not joined:
        logger.info(
            "PDF extracted zero characters of text (pages=%s). Likely scanned/image-only or empty.",
            n_pages,
        )
    return joined


def extract_text_docx(data: bytes) -> str:
    """Extract plain text from DOCX bytes."""
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover
        raise CVExtractionError(
            "DOCX text extraction needs the 'python-docx' package. "
            "Run: pip install python-docx  (or pip install -r requirements.txt)."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise CVExtractionError("Could not read this DOCX (file may be corrupted).") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Tables often hold skills / dates
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs).strip()


def extract_text_from_cv_bytes(*, filename: str, data: bytes, max_bytes: int) -> str:
    """
    Validate and extract text from a CV file.

    Raises:
        CVFileValidationError, CVExtractionError
    """
    ext = validate_cv_upload(filename=filename, data=data, max_bytes=max_bytes)
    logger.info(
        "CV extract start: filename=%s ext=%s size_bytes=%s max_bytes=%s",
        filename,
        ext,
        len(data),
        max_bytes,
    )
    if ext == ".pdf":
        text = extract_text_pdf(data)
        logger.info("CV extract done: pdf text_len=%s", len(text))
        return text
    if ext == ".docx":
        text = extract_text_docx(data)
        logger.info("CV extract done: docx text_len=%s", len(text))
        return text
    raise CVFileValidationError("Unsupported file type.")
